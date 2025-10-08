# -*- coding: utf-8 -*-
"""
02_ExtractArticles.py — ToC-indexed splitter (robust; no regex)

Input:
  inputs/Contract01.docx

Output:
  out/target_articles.jsonl

Also writes debug files:
  out/target_toc_parsed.txt
  out/target_body_preview.txt
  out/target_title_positions.txt
"""

import json
from pathlib import Path
from docx import Document

IN_DIR  = Path("inputs")
OUT_DIR = Path("out"); OUT_DIR.mkdir(parents=True, exist_ok=True)
TARGET  = IN_DIR / "Contract01.docx"

# ---------- utils ----------
SPACE_SET = {" ", "\t", "\u3000", "\u00A0"}
PUNC_DROP = set("：:、，,。.;；．.-—─()（）[]【】<>〈〉《》“”\"'’·•‧")
DOTLIKE   = set(".．・･·‧•●○∙⋅․‥…︙")

def strip_spaces(s: str) -> str:
    return "".join(ch for ch in s if ch not in SPACE_SET) if s else ""

def norm(s: str) -> str:
    """Remove spaces & common punctuation for robust matching."""
    return "".join(ch for ch in (s or "") if ch not in SPACE_SET and ch not in PUNC_DROP)

CN_NUM  = {'零':0,'〇':0,'一':1,'二':2,'兩':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9}
CN_UNIT = {'十':10,'百':100,'千':1000,'萬':10000}
def cn2int(s: str):
    if not s: return None
    s = strip_spaces(s)
    total=0; num=0; used=False
    for ch in s:
        if ch in CN_NUM: num = CN_NUM[ch]
        elif ch in CN_UNIT:
            u = CN_UNIT[ch]
            if num==0: num=1
            total += num*u; num=0; used=True
        else: return None
    total += num
    return total if (used or total!=0 or s in ("零","〇")) else None

# ---------- IO ----------
def read_docx_lines(path: Path):
    doc = Document(str(path))
    lines=[]
    for p in doc.paragraphs:
        t=(p.text or "").rstrip()
        if t.strip(): lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t=(para.text or "").rstrip()
                    if t.strip(): lines.append(t)
    return lines

# ---------- ToC parsing ----------
def parse_toc(lines):
    """
    Parse ToC lines near the top: '第…條 <title> ... <page>'
    Returns (toc, body_start_idx)
      toc = list of dicts {no_int,no_cn,title,title_norm}
    """
    toc=[]; seen=False; body_idx=0
    for i, raw in enumerate(lines[:300]):  # ToC usually near top
        s = raw.strip()
        if ("第" in s) and ("條" in s):
            # heuristic: ToC-like line?
            dot_count = sum(1 for ch in s if ch in DOTLIKE)
            has_tab = ("\t" in s)
            tokens = s.split()
            tail_is_page = (tokens and tokens[-1].isdigit() and len(tokens[-1]) <= 3)
            if not (dot_count >= 3 or has_tab or tail_is_page or s == "目錄"):
                if seen:
                    body_idx = i
                    break
                else:
                    continue
            # extract number between 第 and 條
            try:
                i_d = s.index("第"); i_t = s.index("條", i_d+1)
            except ValueError:
                continue
            no_cn = strip_spaces(s[i_d+1:i_t])
            no_int = cn2int(no_cn)
            if no_int is None:
                continue
            # rest after 條 -> title (trim leading seps; drop trailing page/dots)
            rest = s[i_t+1:].lstrip()
            while rest and rest[0] in ("：", ":", "、", "．", ".", "-", "—", "─", " ", "\t", "\u3000"):
                rest = rest[1:]
            parts = rest.split()
            if parts and parts[-1].isdigit() and len(parts[-1]) <= 3:
                parts = parts[:-1]
            title = " ".join(parts).strip()
            while title and title[-1] in DOTLIKE:
                title = title[:-1].rstrip()
            title = title.strip()
            if not title or "附件" in title:
                continue
            toc.append({"no_int": no_int, "no_cn": no_cn, "title": title, "title_norm": norm(title)})
            seen=True
            body_idx = i+1
        else:
            if seen:
                body_idx = i
                break
    toc.sort(key=lambda e: e["no_int"])
    return toc, body_idx

def is_attachment_heading(s: str) -> bool:
    return s.strip().startswith("附件")

# ---------- heading locator (strict + relaxed) ----------
def find_title_positions(body_lines, toc, max_wrap=8, max_offset=4):
    """
    Locate each ToC title in the body.
    STRICT: joined window must START with title (offset <= max_offset).
    RELAXED fallback: title may appear anywhere in joined window (still compute accurate span).
    Returns list of dicts: {no_int,no_cn,title,start_idx,span}
    """
    body_norm = [norm(x) for x in body_lines]
    out=[]
    last_idx = 0
    L = len(body_lines)

    def joined_norm(i, k):
        s = ""
        for t in range(k):
            if i+t < L:
                s += body_norm[i+t]
        return s

    def compute_span(i, pos_in_join, title_len, max_k):
        """
        Minimal number of lines from i to cover the title up to pos+len(title).
        """
        need = pos_in_join + title_len
        got = 0
        span = 0
        for t in range(max_k):
            if i+t < L:
                got += len(body_norm[i+t])
                span += 1
                if got >= need:
                    break
        return max(1, span)

    for e in toc:
        tnorm = e["title_norm"]
        found = None
        found_span = 1
        scan_from = last_idx

        # STRICT
        for i in range(scan_from, L):
            if not body_norm[i] or is_attachment_heading(body_lines[i]):
                continue
            for k in range(1, max_wrap+1):
                cand = joined_norm(i, k)
                if not cand: 
                    continue
                pos = cand.find(tnorm) if tnorm else -1
                if tnorm and pos != -1 and pos <= max_offset:
                    found = i
                    found_span = compute_span(i, pos, len(tnorm), k)
                    break
            if found is not None:
                break

        # RELAXED fallback
        if found is None:
            for i in range(scan_from, L):
                if not body_norm[i] or is_attachment_heading(body_lines[i]):
                    continue
                for k in range(1, max_wrap+1):
                    cand = joined_norm(i, k)
                    if not cand:
                        continue
                    pos = cand.find(tnorm) if tnorm else -1
                    if tnorm and pos != -1:
                        found = i
                        found_span = compute_span(i, pos, len(tnorm), k)
                        break
                if found is not None:
                    break

        out.append({
            "no_int": e["no_int"],
            "no_cn": e["no_cn"],
            "title": e["title"],
            "start_idx": found,
            "span": found_span if found is not None else 0
        })
        if found is not None:
            last_idx = found + found_span  # move past this heading

    # sort by start index (None at end)
    out.sort(key=lambda d: (d["start_idx"] is None, d["start_idx"] if d["start_idx"] is not None else 10**9))
    return out

# ---------- slicing ----------
def slice_blocks(body_lines, positions):
    """
    Cut body into blocks: for each heading, content is from (start+span) to next start.
    """
    pos = [p for p in positions if p["start_idx"] is not None]
    blocks=[]
    for idx, p in enumerate(pos):
        start = p["start_idx"]
        span  = p["span"] or 1
        content_start = start + span
        content_end   = pos[idx+1]["start_idx"] if idx+1 < len(pos) else len(body_lines)

        buf=[]
        for i in range(content_start, max(content_start, content_end)):
            s = body_lines[i].strip()
            if is_attachment_heading(s):
                continue
            buf.append(body_lines[i])

        blocks.append({
            "article_no": p["no_int"],
            "article_no_cn": p["no_cn"],
            "title": p["title"],
            "text": "\n".join(buf).strip()
        })
    return blocks

# ---------- main ----------
def main():
    if not TARGET.exists():
        raise SystemExit(f"❌ Not found: {TARGET}")

    # read lines
    lines = read_docx_lines(TARGET)

    # parse ToC and compute where body starts
    toc, body_idx = parse_toc(lines)
    if not toc:
        raise SystemExit("❌ Could not parse ToC from the document.")
    (OUT_DIR / "target_toc_parsed.txt").write_text(
        "\n".join([f"{e['no_cn']}({e['no_int']}) {e['title']}" for e in toc]),
        encoding="utf-8"
    )

    # body lines
    body = [ln for ln in lines[body_idx:] if ln and ln.strip()]
    (OUT_DIR / "target_body_preview.txt").write_text("\n".join(body[:300]), encoding="utf-8")

    # locate headings with strict+relaxed search
    positions = find_title_positions(body, toc, max_wrap=8, max_offset=4)
    with open(OUT_DIR / "target_title_positions.txt", "w", encoding="utf-8") as f:
        for p in positions:
            f.write(f"{p['no_cn']:>4}({p['no_int']:>2})  start={p['start_idx'] if p['start_idx'] is not None else 'NOT FOUND'}  span={p['span']}  | {p['title']}\n")

    # slice into article blocks
    blocks = slice_blocks(body, positions)

    # write JSONL
    out_jsonl = OUT_DIR / "target_articles.jsonl"
    with out_jsonl.open("w", encoding="utf-8") as f:
        for b in blocks:
            f.write(json.dumps(b, ensure_ascii=False) + "\n")

    print(f"Parsed {len(blocks)} articles → {out_jsonl}")
    if blocks:
        print("Range:", blocks[0]["article_no"], "…", blocks[-1]["article_no"])
        print("Check out/target_title_positions.txt for heading start & span of each article.")

if __name__ == "__main__":
    main()
