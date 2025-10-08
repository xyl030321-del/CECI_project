# -*- coding: utf-8 -*-
"""
diag_match.py — inspect ToC vs Body matching (no edits to your files)
Creates:
  out/diag_toc.txt         -> parsed ToC entries
  out/diag_body.txt        -> body lines with line numbers
  out/diag_matches.txt     -> which ToC titles were matched to which line
"""

from pathlib import Path
from docx import Document

IN = Path("inputs") / "Contract01.docx"
OUT = Path("out"); OUT.mkdir(exist_ok=True, parents=True)

SPACE_SET = {" ", "\t", "\u3000", "\u00A0"}
PUNC_DROP = set("：:、，,。.;；．.-—─()（）[]【】<>〈〉《》“”\"'’·•‧")
DOTLIKE   = set(".．・･·‧•●○∙⋅․‥…︙")

def strip_spaces(s):
    return "".join(ch for ch in s if ch not in SPACE_SET) if s else ""

def norm(s):
    return "".join(ch for ch in (s or "") if ch not in SPACE_SET and ch not in PUNC_DROP)

# cn numerals -> int
CN_NUM  = {'零':0,'〇':0,'一':1,'二':2,'兩':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9}
CN_UNIT = {'十':10,'百':100,'千':1000,'萬':10000}
def cn2int(s):
    if not s: return None
    s = strip_spaces(s)
    total=0; num=0; used=False
    for ch in s:
        if ch in CN_NUM: num = CN_NUM[ch]
        elif ch in CN_UNIT:
            u = CN_UNIT[ch]
            if num==0: num=1
            total += num*u; num=0; used=True
        else: return None
    total += num
    return total if (used or total!=0 or s in ("零","〇")) else None

def read_docx_lines(p: Path):
    doc = Document(str(p))
    lines=[]
    for para in doc.paragraphs:
        t=(para.text or "").rstrip()
        if t.strip(): lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t=(para.text or "").rstrip()
                    if t.strip(): lines.append(t)
    return lines

def toc_entry_from_line(line):
    s = (line or "").strip()
    if "第" not in s or "條" not in s: return None
    # detect ToC style: dots/tab/page number
    dot_count = sum(1 for ch in s if ch in DOTLIKE)
    has_tab = "\t" in s
    tokens = s.split()
    tail_is_pageno = tokens and tokens[-1].isdigit() and len(tokens[-1])<=3
    if not (dot_count>=3 or has_tab or tail_is_pageno or s=="目錄"):
        return None
    try:
        i = s.index("第"); j = s.index("條", i+1)
    except ValueError:
        return None
    no_cn = strip_spaces(s[i+1:j]); no_int = cn2int(no_cn)
    if no_int is None: return None
    rest = s[j+1:].lstrip()
    while rest and rest[0] in ("：", ":", "、", "．", ".", "-", "—", "─", " ", "\t", "\u3000"):
        rest = rest[1:]
    parts = rest.split()
    if parts and parts[-1].isdigit() and len(parts[-1])<=3: parts = parts[:-1]
    title = " ".join(parts).strip()
    while title and title[-1] in DOTLIKE: title = title[:-1].rstrip()
    title = title.strip()
    if not title or "附件" in title: return None
    return {"no_int": no_int, "no_cn": no_cn, "title": title, "title_norm": norm(title)}

def parse_toc(lines):
    toc=[]; seen=False; body_idx=0
    for i, ln in enumerate(lines[:300]):
        e = toc_entry_from_line(ln)
        if e:
            toc.append(e); seen=True; body_idx=i+1
        else:
            if seen: body_idx=i; break
    toc.sort(key=lambda x: x["no_int"])
    return toc, body_idx

def main():
    if not IN.exists():
        raise SystemExit(f"❌ Not found: {IN}")
    lines = read_docx_lines(IN)

    # parse ToC
    toc, body_idx = parse_toc(lines)
    (OUT/"diag_toc.txt").write_text(
        "\n".join([f"{e['no_cn']}({e['no_int']}) {e['title']}" for e in toc]),
        encoding="utf-8"
    )

    # write body with line numbers
    body = [ln for ln in lines[body_idx:] if ln and ln.strip()]
    (OUT/"diag_body.txt").write_text(
        "\n".join([f"{i+1:04d}: {ln}" for i,ln in enumerate(body)]),
        encoding="utf-8"
    )

    # try to match each ToC title to a body line (within 3-line window, normalized)
    matches=[]
    body_norm = [norm(ln) for ln in body]
    for e in toc:
        tnorm = e["title_norm"]
        found = None
        for i in range(len(body)):
            # join up to 3 lines to handle wrapped headings
            join1 = body_norm[i]
            join2 = (join1 + (body_norm[i+1] if i+1<len(body) else ""))
            join3 = (join2 + (body_norm[i+2] if i+2<len(body) else ""))
            if (tnorm and (join1.startswith(tnorm) or join2.startswith(tnorm) or join3.startswith(tnorm))):
                found = i+1  # 1-based in diag
                break
        matches.append((e["no_int"], e["no_cn"], e["title"], found))

    # save matches
    with open(OUT/"diag_matches.txt", "w", encoding="utf-8") as f:
        for no_int, no_cn, title, idx in matches:
            f.write(f"{no_cn:>4}({no_int:>2})  -> body line {idx if idx else 'NOT FOUND'}  | {title}\n")

    print("✅ Wrote:")
    print(" - out/diag_toc.txt")
    print(" - out/diag_body.txt")
    print(" - out/diag_matches.txt")
    nf = [m for m in matches if not m[3]]
    if nf:
        print(f"⚠️ {len(nf)} titles not found in body. Open out/diag_matches.txt to see which.")
    else:
        print("🎯 All ToC titles were located in the body.")
        
if __name__ == "__main__":
    main()
